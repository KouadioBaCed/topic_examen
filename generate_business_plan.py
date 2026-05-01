"""Generate a Word business plan for the Fin Mark / BIFA Institute platform."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x0A, 0x4A, 0x8F)   # bleu institutionnel
ACCENT = RGBColor(0x06, 0xB6, 0xD4)    # cyan
DARK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x55, 0x65, 0x75)


def set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=DARK, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = DARK
        r1.font.name = "Calibri"
        r2 = p.add_run(" " + text)
        r2.font.size = Pt(11)
        r2.font.color.rgb = DARK
        r2.font.name = "Calibri"
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = DARK
        r.font.name = "Calibri"
    return p


def add_kv_table(doc, rows, col_widths=(Cm(5), Cm(11))):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        c1 = table.cell(i, 0)
        c2 = table.cell(i, 1)
        c1.width = col_widths[0]
        c2.width = col_widths[1]
        set_cell_bg(c1, "F1F5F9")
        c1.paragraphs[0].text = ""
        c2.paragraphs[0].text = ""
        r1 = c1.paragraphs[0].add_run(k)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = PRIMARY
        r1.font.name = "Calibri"
        r2 = c2.paragraphs[0].add_run(v)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = DARK
        r2.font.name = "Calibri"
    return table


def add_data_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        set_cell_bg(hdr[i], "0A4A8F")
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Calibri"
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = DARK
            r.font.name = "Calibri"
    return table


def add_section_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break()


def main(output_path: str):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # --------- COUVERTURE ---------
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("BUSINESS PLAN")
    r.bold = True
    r.font.size = Pt(34)
    r.font.color.rgb = PRIMARY
    r.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Fin Mark — Plateforme de préparation aux certifications internationales")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = ACCENT
    r.font.name = "Calibri"

    by = doc.add_paragraph()
    by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = by.add_run("Une initiative de BIFA Institute")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = GREY

    add_para(doc, "")
    add_para(doc, "")

    add_kv_table(doc, [
        ("Porteur du projet", "BIFA Institute"),
        ("Marque produit", "Fin Mark"),
        ("Siège social", "Abidjan, Plateau, Immeuble Arc-en-Ciel — Côte d'Ivoire"),
        ("Site web institutionnel", "https://bifainstitute.netlify.app"),
        ("Contact commercial", "programmes@bifainstitute.com  |  +225 07 00 74 82 28"),
        ("Statut du document", "Business plan — version exécutive"),
        ("Date d'édition", "Avril 2026"),
        ("Confidentialité", "Document confidentiel — diffusion restreinte"),
    ])

    doc.add_page_break()

    # --------- SOMMAIRE ---------
    add_heading(doc, "Sommaire", level=1)
    sommaire = [
        "1. Résumé exécutif",
        "2. Présentation de l'entreprise et du produit",
        "3. Analyse du marché",
        "4. Positionnement stratégique et proposition de valeur",
        "5. Modèle économique",
        "6. Stratégie commerciale et marketing",
        "7. Architecture produit et technologie",
        "8. Organisation et équipe",
        "9. Plan opérationnel et feuille de route",
        "10. Prévisions financières",
        "11. Analyse des risques",
        "12. Demande de financement et utilisation des fonds",
        "13. Annexes",
    ]
    for s in sommaire:
        add_para(doc, s, size=11, color=DARK)

    doc.add_page_break()

    # --------- 1. RÉSUMÉ EXÉCUTIF ---------
    add_heading(doc, "1. Résumé exécutif", level=1)
    add_para(doc,
        "Fin Mark est la plateforme SaaS de préparation aux certifications professionnelles "
        "internationales développée par BIFA Institute. La solution adresse une problématique forte "
        "en Afrique francophone : l'accès à une préparation structurée, abordable et de qualité aux "
        "examens internationaux exigeants (PMP®, CAPM®, CIA, CISA®, CFE, PL-300, MOS, PM4NGOs).",
    )
    add_para(doc,
        "La plateforme propose plus de 4 500 questions à choix multiples avec corrections "
        "détaillées, des examens blancs en conditions réelles, un suivi de progression et une "
        "révision structurée des terminologies, le tout en interface bilingue français / anglais.",
    )
    add_heading(doc, "Faits saillants", level=2)
    add_bullet(doc, "10+ certifications internationales couvertes en un seul abonnement.", "Couverture :")
    add_bullet(doc, "Pédagogie validée par plus de 97 % de réussite déclarée par les certifiés interrogés.", "Performance :")
    add_bullet(doc, "Un marché adressable francophone large et sous-servi (Afrique de l'Ouest, Afrique centrale, diaspora).", "Marché :")
    add_bullet(doc, "Modèle SaaS récurrent, marges logicielles, scalabilité naturelle vers de nouvelles certifications.", "Modèle :")
    add_bullet(doc, "Une recherche de financement de 120 000 USD pour accélérer acquisition, contenu et expansion régionale.", "Ambition :")

    doc.add_page_break()

    # --------- 2. ENTREPRISE & PRODUIT ---------
    add_heading(doc, "2. Présentation de l'entreprise et du produit", level=1)

    add_heading(doc, "2.1 BIFA Institute, le porteur du projet", level=2)
    add_para(doc,
        "BIFA Institute est un institut de formation professionnelle basé à Abidjan, spécialisé "
        "dans la formation et l'accompagnement aux certifications internationales en gestion de "
        "projet, audit, conformité, finance et data. L'institut capitalise sur un réseau de "
        "formateurs certifiés, des partenariats avec les organismes émetteurs (PMI, IIA, ACFE, "
        "Microsoft, PM4NGOs) et un historique de candidats certifiés à fort taux de réussite.",
    )

    add_heading(doc, "2.2 Fin Mark, la solution numérique", level=2)
    add_para(doc,
        "Fin Mark est l'extension digitale de l'offre de BIFA Institute. La plateforme web "
        "transforme l'accompagnement présentiel en un produit SaaS accessible 24/7, scalable et "
        "exportable au-delà des frontières ivoiriennes.",
    )
    add_para(doc, "Fonctionnalités clés :", bold=True)
    add_bullet(doc, "Banque de questions QCM (choix unique et multiples) avec explications détaillées.")
    add_bullet(doc, "Mode Examen chronométré reproduisant les conditions réelles de chaque certification.")
    add_bullet(doc, "Mini-examens et examens blancs par domaine et par niveau.")
    add_bullet(doc, "Révision structurée des terminologies, formules et concepts clés.")
    add_bullet(doc, "Tableau de bord personnel, historique des résultats, analyse fine des performances.")
    add_bullet(doc, "Espace administrateur pour la gestion des accès, des cohortes et de la facturation.")
    add_bullet(doc, "Bilinguisme français / anglais natif sur l'ensemble de l'expérience.")

    add_heading(doc, "2.3 Catalogue de certifications couvertes", level=2)
    add_data_table(doc,
        headers=["Catégorie", "Certifications", "Public cible"],
        rows=[
            ("Gestion de projets", "PMP®, CAPM®, PM4NGOs (Project DPro)", "Chefs de projet, PMO, consultants"),
            ("Programmes & ONG", "Program DPro, Social Good DPro, MEAL DPro, Finance DPro", "Cadres ONG, bailleurs, coopération"),
            ("Audit & conformité", "CIA Parts 1/2/3, CISA®, CFE", "Auditeurs, contrôleurs, risk managers"),
            ("Data & reporting", "PL-300 (Power BI), MOS Excel, MOS Word", "Analystes, data professionals, consultants"),
        ],
    )

    doc.add_page_break()

    # --------- 3. MARCHÉ ---------
    add_heading(doc, "3. Analyse du marché", level=1)

    add_heading(doc, "3.1 Tendances structurelles", level=2)
    add_bullet(doc, "Demande croissante de certifications internationales pour accéder aux postes à responsabilités dans les multinationales, ONG et grandes administrations africaines.")
    add_bullet(doc, "Pénétration rapide d'Internet et du paiement mobile en Afrique francophone, créant un terreau favorable au e-learning.")
    add_bullet(doc, "Coût de l'examen en devise (200 à 700 USD selon la certification) : les candidats cherchent à maximiser leurs chances dès la première tentative.")
    add_bullet(doc, "Pénurie d'offres de préparation locales structurées et abordables, surtout en français.")

    add_heading(doc, "3.2 Marché adressable", level=2)
    add_data_table(doc,
        headers=["Indicateur", "Estimation", "Source"],
        rows=[
            ("TAM — francophones intéressés par une certif. pro / an", "≈ 1 500 000", "Estimation interne, CIO Mag, IFC"),
            ("SAM — Afrique de l'Ouest francophone numérique", "≈ 250 000", "Estimation interne"),
            ("SOM à 3 ans (part atteignable)", "≈ 8 000 utilisateurs payants", "Hypothèses BIFA Institute"),
            ("Panier moyen annuel cible", "≈ 90 USD / utilisateur", "Benchmark concurrents"),
        ],
    )

    add_heading(doc, "3.3 Cibles prioritaires", level=2)
    add_bullet(doc, "Jeunes professionnels 25-40 ans cherchant une montée en compétence et une mobilité internationale.", "Persona 1 :")
    add_bullet(doc, "Cadres d'ONG et bailleurs internationaux ayant besoin de PM4NGOs / MEAL / Finance DPro.", "Persona 2 :")
    add_bullet(doc, "Auditeurs internes, contrôleurs et risk managers en banque, télécoms et secteur public.", "Persona 3 :")
    add_bullet(doc, "Entreprises et institutions souhaitant former massivement leurs équipes (B2B).", "Persona 4 :")

    add_heading(doc, "3.4 Concurrence", level=2)
    add_data_table(doc,
        headers=["Acteur", "Forces", "Faiblesses vs Fin Mark"],
        rows=[
            ("PMI Study Hall, Pocket Prep", "Notoriété internationale, contenu PMP", "Onéreux, anglais uniquement, pas de support local"),
            ("Udemy / Coursera", "Catalogue large, prix bas", "Pédagogie inégale, pas d'examens blancs ciblés Afrique"),
            ("Cabinets locaux de formation", "Présentiel, réseau local", "Pas de plateforme digitale, capacité limitée"),
            ("Banques de questions piratées", "Gratuites", "Illégales, contenu obsolète, aucune garantie"),
        ],
    )

    doc.add_page_break()

    # --------- 4. POSITIONNEMENT ---------
    add_heading(doc, "4. Positionnement stratégique et proposition de valeur", level=1)
    add_para(doc, "Promesse de marque :", bold=True)
    add_para(doc,
        "« Réussir les certifications internationales avec une méthode pensée pour les "
        "professionnels francophones. »", italic=True,
    )

    add_heading(doc, "4.1 Différenciateurs", level=2)
    add_bullet(doc, "Couverture multi-certifications dans un seul abonnement (vs un produit = une certif).")
    add_bullet(doc, "Pédagogie en français, alignée sur les référentiels officiels les plus récents.")
    add_bullet(doc, "Questions très proches de l'examen réel (94 % de proximité déclarée par les certifiés).")
    add_bullet(doc, "Encadrement humain par des formateurs BIFA Institute (mentorat, sessions de coaching).")
    add_bullet(doc, "Paiement local : Mobile Money, carte, virement — adapté au contexte africain.")

    add_heading(doc, "4.2 Avantage compétitif durable", level=2)
    add_bullet(doc, "Capital pédagogique : milliers de questions originales rédigées et validées par des certifiés.")
    add_bullet(doc, "Communauté : réseau d'alumni qui alimente le bouche-à-oreille et le contenu.")
    add_bullet(doc, "Marque BIFA Institute : crédibilité institutionnelle déjà installée en Côte d'Ivoire.")
    add_bullet(doc, "Effet de levier B2B : entreprises et ONG signent des contrats annuels pour leurs équipes.")

    doc.add_page_break()

    # --------- 5. MODÈLE ÉCONOMIQUE ---------
    add_heading(doc, "5. Modèle économique", level=1)

    add_heading(doc, "5.1 Sources de revenus", level=2)
    add_bullet(doc, "Abonnements individuels mensuels et annuels (récurrent).", "B2C SaaS :")
    add_bullet(doc, "Licences groupées pour entreprises, ONG et écoles supérieures.", "B2B :")
    add_bullet(doc, "Coaching, classes virtuelles, bootcamps (premium add-on).", "Services :")
    add_bullet(doc, "Cours certifiants, examens blancs sponsorisés, cobranding avec organismes émetteurs.", "Partenariats :")

    add_heading(doc, "5.2 Grille tarifaire indicative", level=2)
    add_data_table(doc,
        headers=["Offre", "Prix", "Inclus"],
        rows=[
            ("Découverte (7 jours)", "Gratuit", "1 mini-examen + accès limité aux QCM"),
            ("Mensuel", "15 000 FCFA / mois", "1 certification au choix, accès illimité"),
            ("Trimestriel", "35 000 FCFA / 3 mois", "1 certification, suivi mentor + simulation"),
            ("Annuel Premium", "120 000 FCFA / an", "Toutes les certifications + coaching"),
            ("B2B (par siège)", "Sur devis", "Pack 10 à 500 sièges, tableau de bord RH"),
        ],
    )

    add_heading(doc, "5.3 Économie unitaire cible (en année 2)", level=2)
    add_data_table(doc,
        headers=["Indicateur", "Valeur cible"],
        rows=[
            ("Coût d'acquisition client (CAC)", "≈ 18 USD"),
            ("Revenu moyen par utilisateur (ARPU annualisé)", "≈ 90 USD"),
            ("Marge brute SaaS", "≈ 78 %"),
            ("Lifetime Value (LTV)", "≈ 160 USD"),
            ("Ratio LTV / CAC", "≈ 8,9 x"),
            ("Churn mensuel cible", "< 5 %"),
        ],
    )

    doc.add_page_break()

    # --------- 6. STRATÉGIE GO-TO-MARKET ---------
    add_heading(doc, "6. Stratégie commerciale et marketing", level=1)

    add_heading(doc, "6.1 Acquisition", level=2)
    add_bullet(doc, "SEO et contenus longs : guides PMP, CIA, CFE en français, optimisés pour la recherche.")
    add_bullet(doc, "Réseaux sociaux : LinkedIn (B2B et professionnels), Facebook (volume), YouTube (tutoriels et témoignages).")
    add_bullet(doc, "Webinaires et masterclasses gratuits comme aimants à inscriptions.")
    add_bullet(doc, "Partenariats avec ONG, cabinets de conseil, grandes écoles et universités.")
    add_bullet(doc, "Programme d'ambassadeurs alumni avec commissions sur abonnements.")

    add_heading(doc, "6.2 Conversion et rétention", level=2)
    add_bullet(doc, "Essai gratuit de 7 jours sans carte bancaire.")
    add_bullet(doc, "Onboarding guidé : test de positionnement, plan d'étude personnalisé.")
    add_bullet(doc, "Notifications de progression, rappels d'examen, gamification (badges, classements).")
    add_bullet(doc, "Programme « Réussite garantie » : extension gratuite si l'examen est échoué malgré un parcours complet.")

    add_heading(doc, "6.3 Stratégie B2B", level=2)
    add_bullet(doc, "Démarche commerciale ciblée vers les directions RH et formation.")
    add_bullet(doc, "Tableaux de bord RH : suivi de progression et de certification des collaborateurs.")
    add_bullet(doc, "Contrats annuels avec engagement et facturation centralisée.")

    doc.add_page_break()

    # --------- 7. PRODUIT & TECHNO ---------
    add_heading(doc, "7. Architecture produit et technologie", level=1)

    add_heading(doc, "7.1 Stack technique", level=2)
    add_bullet(doc, "React 19 + TypeScript + Vite + TailwindCSS, déploiement Netlify.", "Frontend :")
    add_bullet(doc, "Firebase (Authentication, Firestore, Hosting, Cloud Functions).", "Backend :")
    add_bullet(doc, "Mobile Money via agrégateur local, Stripe pour la carte internationale.", "Paiements :")
    add_bullet(doc, "Suivi anonymisé via Google Analytics 4, dashboards internes.", "Analytics :")
    add_bullet(doc, "Authentification sécurisée, RBAC admin / utilisateur, conformité RGPD.", "Sécurité :")

    add_heading(doc, "7.2 Roadmap produit (12 mois)", level=2)
    add_data_table(doc,
        headers=["Trimestre", "Livrables clés"],
        rows=[
            ("T1", "Refonte tarification, paiement Mobile Money, gestion des cohortes B2B"),
            ("T2", "Application mobile (PWA), mode hors ligne, mentorat intégré"),
            ("T3", "Lancement de 3 nouvelles certifications (Scrum, ITIL, ISO 27001)"),
            ("T4", "IA pédagogique : explications adaptatives, feedback personnalisé"),
        ],
    )

    add_heading(doc, "7.3 Indicateurs produit suivis", level=2)
    add_bullet(doc, "Taux d'activation à 7 jours, taux de complétion d'examen blanc, NPS, churn par cohorte, taux de réussite déclaré post-certification.")

    doc.add_page_break()

    # --------- 8. ÉQUIPE ---------
    add_heading(doc, "8. Organisation et équipe", level=1)

    add_para(doc,
        "Le projet s'appuie sur l'équipe de BIFA Institute et un noyau produit dédié. "
        "L'organisation cible à 18 mois est la suivante :",
    )

    add_data_table(doc,
        headers=["Rôle", "Responsabilités principales", "Statut"],
        rows=[
            ("Direction générale", "Vision, partenariats, levée de fonds", "En place"),
            ("Direction pédagogique", "Catalogue, qualité des questions, formateurs", "En place"),
            ("Lead produit / tech", "Roadmap, architecture, sécurité", "À renforcer"),
            ("Développeurs full-stack (2)", "Frontend React, Firebase, intégrations", "1 en place, 1 à recruter"),
            ("Growth / marketing", "Acquisition, contenu, social media", "À recruter"),
            ("Customer success", "Onboarding, support, suivi B2B", "À recruter"),
            ("Réseau de formateurs certifiés", "Création de contenu, coaching", "Pool de 12 contributeurs"),
        ],
    )

    doc.add_page_break()

    # --------- 9. PLAN OPÉRATIONNEL ---------
    add_heading(doc, "9. Plan opérationnel et feuille de route", level=1)

    add_data_table(doc,
        headers=["Phase", "Période", "Objectifs"],
        rows=[
            ("Phase 1 – Consolidation", "T2 2026", "500 abonnés actifs, paiement Mobile Money en production, NPS > 50"),
            ("Phase 2 – Expansion", "T3-T4 2026", "2 000 abonnés, 5 contrats B2B, ouverture Sénégal et Cameroun"),
            ("Phase 3 – Scale", "2027", "8 000 abonnés, 20 contrats B2B, app mobile, 3 nouvelles certifications"),
            ("Phase 4 – Régionalisation", "2028", "Présence dans 8 pays francophones, 25 000 abonnés actifs"),
        ],
    )

    doc.add_page_break()

    # --------- 10. FINANCES ---------
    add_heading(doc, "10. Prévisions financières", level=1)

    add_heading(doc, "10.1 Hypothèses", level=2)
    add_bullet(doc, "Panier moyen annuel : 90 USD (mix B2C / B2B).")
    add_bullet(doc, "Coût d'acquisition moyen : 18 USD, en baisse avec la maturité de la marque.")
    add_bullet(doc, "Marge brute SaaS : 78 % en année 1, 82 % en année 3.")
    add_bullet(doc, "Croissance organique soutenue par bouche-à-oreille (effet alumni).")

    add_heading(doc, "10.2 Compte de résultat prévisionnel (en USD)", level=2)
    add_data_table(doc,
        headers=["Poste", "Année 1", "Année 2", "Année 3"],
        rows=[
            ("Utilisateurs payants en fin d'année", "1 200", "3 800", "8 000"),
            ("Chiffre d'affaires", "85 000", "260 000", "640 000"),
            ("Coûts d'infrastructure et licences", "10 000", "26 000", "55 000"),
            ("Coûts de contenu et formateurs", "18 000", "40 000", "75 000"),
            ("Marketing et acquisition", "25 000", "60 000", "120 000"),
            ("Salaires et charges", "55 000", "120 000", "210 000"),
            ("Frais généraux", "12 000", "20 000", "30 000"),
            ("Total charges", "120 000", "266 000", "490 000"),
            ("Résultat d'exploitation", "-35 000", "-6 000", "+150 000"),
        ],
    )

    add_heading(doc, "10.3 Trajectoire", level=2)
    add_bullet(doc, "Année 1 : phase d'investissement, point mort opérationnel non atteint, cash brûlé maîtrisé.")
    add_bullet(doc, "Année 2 : quasi-équilibre, mise à l'échelle de la base d'abonnés.")
    add_bullet(doc, "Année 3 : profitabilité opérationnelle, marge nette > 20 %, capacité d'autofinancement.")

    doc.add_page_break()

    # --------- 11. RISQUES ---------
    add_heading(doc, "11. Analyse des risques", level=1)
    add_data_table(doc,
        headers=["Risque", "Impact", "Mesure d'atténuation"],
        rows=[
            ("Piratage du contenu pédagogique", "Élevé", "Filigranage, sessions par utilisateur, watermarking dynamique"),
            ("Évolution des référentiels (PMI, IIA)", "Moyen", "Veille, équipe pédagogique, contrats avec experts certifiés"),
            ("Concurrence internationale (Pocket Prep, etc.)", "Moyen", "Différenciation francophone et ancrage local"),
            ("Dépendance à Firebase / Netlify", "Moyen", "Architecture portable, exports réguliers, plan B Supabase"),
            ("Risque de paiement (Mobile Money)", "Moyen", "Multi-agrégateurs, réconciliation automatisée"),
            ("Fluctuation FCFA / USD", "Faible-Moyen", "Tarification locale + facturation USD pour B2B export"),
            ("Recrutement tech en Côte d'Ivoire", "Moyen", "Salaire compétitif, télétravail, recrutement panafricain"),
        ],
    )

    doc.add_page_break()

    # --------- 12. FINANCEMENT ---------
    add_heading(doc, "12. Demande de financement et utilisation des fonds", level=1)
    add_para(doc,
        "BIFA Institute recherche un financement de 120 000 USD sous forme d'amorçage "
        "(equity, dette convertible ou subvention) afin d'accélérer la trajectoire "
        "présentée ci-dessus. Cette enveloppe couvre 18 mois d'opérations et permet "
        "d'atteindre le point mort opérationnel en année 2.",
    )
    add_data_table(doc,
        headers=["Poste", "Montant (USD)", "Part"],
        rows=[
            ("Renforcement produit et tech (2 dev, infra)", "45 000", "37,5 %"),
            ("Acquisition et marketing", "35 000", "29,2 %"),
            ("Production de contenu et certifications additionnelles", "20 000", "16,7 %"),
            ("Structuration commerciale B2B", "12 000", "10,0 %"),
            ("Frais juridiques, comptables et imprévus", "8 000", "6,6 %"),
            ("Total", "120 000", "100 %"),
        ],
    )
    add_para(doc, "Retour attendu pour les investisseurs :", bold=True)
    add_bullet(doc, "Profitabilité opérationnelle en année 3.")
    add_bullet(doc, "Valorisation cible à 5 ans : 4 à 6 millions USD (multiple de 6 à 8 x ARR).")
    add_bullet(doc, "Sortie possible via rachat stratégique par un acteur edtech panafricain ou un cabinet de formation international.")

    doc.add_page_break()

    # --------- 13. ANNEXES ---------
    add_heading(doc, "13. Annexes", level=1)
    add_bullet(doc, "Témoignages détaillés et photos de certifiés (alumni Fin Mark / BIFA Institute).")
    add_bullet(doc, "Captures d'écran de la plateforme (mode examen, dashboard, révision).")
    add_bullet(doc, "Liste exhaustive des certifications couvertes et nombre de questions par module.")
    add_bullet(doc, "Détails techniques : schéma d'architecture, modèle de données Firestore, sécurité.")
    add_bullet(doc, "Modèle financier complet (Excel) — disponible sur demande.")

    add_para(doc, "")
    add_para(doc, "")
    add_para(doc, "— Fin du document —", italic=True, color=GREY,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(output_path)
    print(f"OK -> {output_path}")


if __name__ == "__main__":
    main(r"C:\Users\lenovo\Downloads\topic_examen\topic_exam\Business_Plan_Fin_Mark.docx")
